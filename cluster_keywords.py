import pandas as pd


# 1. Загрузка данных
df = pd.read_excel('/Users/oksanagoncarova/Desktop/статьи лето/русистика статья/articles_metadata.xlsx')

import pandas as pd
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.cluster import KMeans


df['keywords_list'] = (
    df['keywords']
      .fillna('')
      .astype(str)
      .apply(lambda x: [kw.strip().lower() for kw in x.split(',') if kw.strip()])
)

# 2. Векторизация «сырого» текста ключевых слов
vectorizer = CountVectorizer(token_pattern=r'(?u)\b\w+\b')
X = vectorizer.fit_transform(df['keywords'].fillna('').astype(str))

# 3. Кластеризация KMeans на 5 групп
k = 5
model = KMeans(n_clusters=k, random_state=42)
df['cluster'] = model.fit_predict(X)

# 4. Присвоение читаемых названий кластерам
cluster_names = {
    0: '1',
    1: '2',
    2: '3',
    3: '4',
    4: '5'
}
df['group'] = df['cluster'].map(cluster_names)

# 5. «Взрыв» списка ключевых слов в строки для полного списка
exploded = df.explode('keywords_list')[['group', 'keywords_list']] \
             .rename(columns={'keywords_list': 'keyword'})

# 6. Полный список ключевых слов по группам
group_keywords = (
    exploded
    .groupby('group')['keyword']
    .unique()
    .apply(list)
    .reset_index()
)
print("Полный список ключевых слов по группам:")
print(group_keywords.to_string(index=False))

# 7. Получение 10 наиболее частотных ключевых слов в каждой группе
#    Считаем частоты в exploded
freq = (
    exploded
    .groupby(['group', 'keyword'])
    .size()
    .reset_index(name='count')
)
top10 = (
    freq
    .sort_values(['group', 'count'], ascending=[True, False])
    .groupby('group')
    .head(10)
    .groupby('group')['keyword']
    .apply(list)
    .reset_index()
    .rename(columns={'keyword': 'top_10_keywords'})
)
print("\n10 наиболее частых ключевых слов по группам:")
print(top10.to_string(index=False))


output_path = '/Users/oksanagoncarova/Desktop/статьи лето/русистика статья/keywords_analysis.xlsx'

with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
    # Полный список ключевых слов по группам
    group_keywords.to_excel(writer,
                            sheet_name='All_Group_Keywords',
                            index=False)
    # Топ-10 самых частых ключевых слов по группам
    top10.to_excel(writer,
                   sheet_name='Top10_Keywords',
                   index=False)

print(f"Результаты сохранены в файл: {output_path}")\


from __future__ import annotations

import re

from collections import Counter
from pathlib import Path
from typing import Dict, List, Set, Tuple

import docx
import nltk
from nltk import bigrams, word_tokenize
from nltk.corpus import stopwords
from pymorphy3 import MorphAnalyzer
from sklearn.decomposition import NMF
from sklearn.feature_extraction.text import TfidfVectorizer

# ─────────────────────────── NLTK ресурсы ───────────────────────────
try:
    BASE_STOP: Set[str] = set(stopwords.words("russian"))
except LookupError:
    nltk.download("stopwords")
    BASE_STOP = set(stopwords.words("russian"))

try:
    nltk.data.find("tokenizers/punkt")
except LookupError:
    nltk.download("punkt")

# ─────────────────────────── Дополнительные стоп-слова ───────────────────────────
EXTRA_STOP_COMMON: Set[str] = {
    "иностранный",
    "язык",
    "интеллект",
    "искусственный",
    "обучение",
    "русский",
    "технология",
    "рки",
    'ия',
    'который',
    'статья',
    'преподавание',
    'посвятить',
    'исследование',
    'особый',
    'внимание',
    'сравнительный',
    'анализ',


}
STOP_COMMON = BASE_STOP.union(EXTRA_STOP_COMMON)

# ─────────────────────────── Лемматизатор ───────────────────────────
morph = MorphAnalyzer()

# ─────────────────────────── Вспомогательные функции ───────────────────────────

def extract_text_from_cell(cell: docx.table._Cell) -> str:  # type: ignore
    """Рекурсивно извлекает текст из ячейки, включая вложенные таблицы."""
    out: List[str] = []
    for p in cell.paragraphs:
        if p.text:
            out.append(p.text)
    for tbl in cell.tables:
        for row in tbl.rows:
            for nested in row.cells:
                out.append(extract_text_from_cell(nested))
    return "\n".join(out)


def load_columns(doc_path: Path) -> Dict[str, List[str]]:
    """Извлекает столбцы «ключевые слова» и «аннотация» (литературу игнорируем)."""
    doc = docx.Document(doc_path)
    data = {"keywords": [], "annotation": []}
    for table in doc.tables:
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        if {"ключевые слова", "аннотация"}.issubset(header):
            idx_kw = header.index("ключевые слова")
            idx_ann = header.index("аннотация")
            for r in table.rows[1:]:
                data["keywords"].append(extract_text_from_cell(r.cells[idx_kw]))
                data["annotation"].append(extract_text_from_cell(r.cells[idx_ann]))
    return data


def preprocess(text: str, stop_set: Set[str]) -> List[str]:
    clean = re.sub(r"[^А-Яа-яЁё]", " ", text)
    lemmas: List[str] = []
    for tok in word_tokenize(clean):
        if tok.isalpha():
            lemma = morph.parse(tok.lower())[0].normal_form
            if lemma not in stop_set:
                lemmas.append(lemma)
    return lemmas


def corpus_from_list(texts: List[str], stop_set: Set[str]) -> List[str]:
    return [" ".join(preprocess(t, stop_set)) for t in texts]

# ─────────────────────────── Статистика ───────────────────────────

def tfidf_top_terms(texts: List[str], top_n: int = 8):
    vect = TfidfVectorizer()
    tfidf = vect.fit_transform(texts)
    mean_scores = tfidf.mean(axis=0).A1
    terms = vect.get_feature_names_out()
    return sorted(zip(terms, mean_scores), key=lambda x: x[1], reverse=True)[:top_n]


def freq_top_terms(texts: List[str], stop_set: Set[str], top_n: int = 8):
    counter = Counter()
    for t in texts:
        counter.update(preprocess(t, stop_set))
    return counter.most_common(top_n)


def top_bigrams(texts: List[str], stop_set: Set[str], top_n: int = 8):
    counter = Counter()
    for t in texts:
        counter.update(["_".join(bg) for bg in bigrams(preprocess(t, stop_set))])
    return counter.most_common(top_n)

# ─────────────────────────── Темы аннотаций (NMF) ───────────────────────────

def nmf_topics(texts: List[str], n_topics: int = 5, top_words: int = 6) -> List[Tuple[int, List[str]]]:
    vect = TfidfVectorizer(max_df=0.9, min_df=2)
    tfidf = vect.fit_transform(texts)
    nmf = NMF(n_components=n_topics, init="nndsvda", random_state=0)
    nmf.fit(tfidf)
    terms = vect.get_feature_names_out()
    topics: List[Tuple[int, List[str]]] = []
    for idx, comp in enumerate(nmf.components_):
        top_idx = comp.argsort()[::-1][:top_words]
        topics.append((idx + 1, [terms[i] for i in top_idx]))
    return topics

# ─────────────────────────── Главная функция ───────────────────────────

def main(doc_path: Path):
    cols = load_columns(doc_path)
    results: Dict[str, Dict[str, List[Tuple[str, float]]]] = {}
    for col in ("keywords", "annotation"):
        texts = cols[col]
        processed = corpus_from_list(texts, STOP_COMMON)
        results[col] = {
            "tfidf": tfidf_top_terms(processed),
            "freq": freq_top_terms(texts, STOP_COMMON),
            "bigrams": top_bigrams(texts, STOP_COMMON),
        }

    topics = nmf_topics(corpus_from_list(cols["annotation"], STOP_COMMON))

    # ─────────── Вывод ───────────
    for col, res in results.items():
        print(f"\n=== {col.upper()} ===")
        print("TF-IDF (top-8):")
        for w, s in res["tfidf"]:
            print(f"  {w:<20} {s:.4f}")
        print("Most frequent (top-8):")
        for w, c in res["freq"]:
            print(f"  {w:<20} {c}")
        print("Top bigrams (top-8):")
        for bg, c in res["bigrams"]:
            print(f"  {bg:<25} {c}")

    print("\n=== TOPICS FROM ANNOTATIONS (NMF) ===")
    for idx, words in topics:
        print(f"Topic {idx}: {', '.join(words)}")




if __name__ == "__main__":
    # ▼▼▼ Укажите путь к своему .docx‑файлу ▼▼▼
    main(Path("/Users/oksanagoncarova/Desktop/статьи лето/русистика статья/2022.docx"))