# requirements:
# pip install python-docx networkx lxml

#
#
# docx_path = "/Users/oksanagoncarova/Desktop/статьи лето/русистика статья/ARTIFICIAL INTELLIGENCE IN FOREIGN LANGUAGE TEACHING.docx"
#
#
# from docx import Document
# import pandas as pd
# import re
# from pathlib import Path
#
# DOCX = Path("/Users/oksanagoncarova/Desktop/статьи лето/русистика статья/ARTIFICIAL INTELLIGENCE IN FOREIGN LANGUAGE TEACHING.docx")         # имя файла
# NODES_CSV = "nodes.csv"
# EDGES_CSV = "edges.csv"
#
# # --- вспомогательные регэкспы ---
# AUTHOR_LINE_RGX = re.compile(r"[А-ЯЁ][а-яё\-]+\s+[А-ЯЁ]\.")  # Фамилия И.
# CITE_AUTHOR_RGX = re.compile(r"^([А-ЯЁ][а-яё\-]+)\s+[А-ЯЁ]\.")  # первая фамилия в ссылке
#
# doc = Document(DOCX)
# table = doc.tables[0]                # главная таблица
#
# articles = []                        # [{'id','authors':[...], 'cited':[...]}]
# for r_i, row in enumerate(table.rows[1:], start=1):           # пропускаем заголовок
#     art_id = f"ART_{r_i}"
#
#     # ---------- авторы статьи ----------
#     txt0 = row.cells[0].text.splitlines()
#     authors_raw = next((ln for ln in txt0 if AUTHOR_LINE_RGX.search(ln)), "")
#     authors = [a.strip() for a in re.split(r",| и ", authors_raw) if a.strip()]
#
#     # ---------- литература (вложенная таблица) ----------
#     cited = set()
#     lit_cell = row.cells[3]
#     if lit_cell.tables:
#         for nrow in lit_cell.tables[0].rows:
#             # во вложенной таблице первая ячейка – номер, остальное – ссылка
#             ref_text = " ".join(c.text.strip() for c in nrow.cells[1:]).strip()
#             m = CITE_AUTHOR_RGX.match(ref_text)
#             if m:
#                 cited.add(m.group(1))             # только фамилию
#
#     articles.append({"id": art_id, "authors": authors, "cited": list(cited)})
#
# # ---------- узлы ----------
# nodes = []
# for art in articles:
#     nodes.append({"Id": art["id"], "Label": art["id"], "Type": "Article"})
# cited_authors = {a for art in articles for a in art["cited"]}
# for a in cited_authors:
#     nodes.append({"Id": f"AUT_{a}", "Label": a, "Type": "Author"})
#
# # ---------- рёбра статья → автор ----------
# edges = []
# for art in articles:
#     for a in art["cited"]:
#         edges.append({
#             "Source": art["id"],
#             "Target": f"AUT_{a}",
#             "Type": "Directed",
#             "Label": "cites",
#             "Weight": 1
#         })
#
# # ---------- выгрузка ----------
# pd.DataFrame(nodes).to_csv(NODES_CSV, index=False)
# pd.DataFrame(edges).to_csv(EDGES_CSV, index=False)
# print(f"Файлы сохранены: {NODES_CSV}, {EDGES_CSV}")
from docx import Document
import pandas as pd
import re
from pathlib import Path
from collections import defaultdict
from itertools import combinations

DOCX = Path("/Users/oksanagoncarova/Desktop/статьи лето/русистика статья/ARTIFICIAL INTELLIGENCE IN FOREIGN LANGUAGE TEACHING.docx")
NODES_CSV = "nodes.csv"
EDGES_CSV = "edges.csv"

# Улучшенные регулярные выражения
AUTHOR_RGX = re.compile(r"[А-ЯЁ][а-яё\-]+\s+[А-ЯЁ]\.[А-ЯЁ]?\.?")  # Фамилия И.О.
CITE_AUTHOR_RGX = re.compile(r"([А-ЯЁ][а-яё\-]+)\s+[А-ЯЁ]\.")  # Фамилия И.


# Нормализация имен авторов
def normalize_author(author):
    return re.sub(r'\s+', ' ', author.strip()).title()


doc = Document(DOCX)
table = doc.tables[0]

articles = []
coauthor_edges = defaultdict(int)  # (author1, author2) -> count

for r_i, row in enumerate(table.rows[1:], start=1):
    art_id = f"ART_{r_i}"
    cell_text = [line.strip() for line in row.cells[0].text.splitlines() if line.strip()]

    # Извлекаем название статьи (первая непустая строка)
    title = cell_text[0] if cell_text else f"Статья_{r_i}"

    # Извлекаем всех авторов
    authors = []
    for line in cell_text[1:]:
        authors.extend([normalize_author(a) for a in AUTHOR_RGX.findall(line)])

    # Обработка литературы
    cited = set()
    lit_cell = row.cells[3]
    if lit_cell.tables:
        for nrow in lit_cell.tables[0].rows:
            ref_text = " ".join(c.text.strip() for c in nrow.cells[1:])
            # Извлекаем всех авторов из ссылки
            ref_authors = {normalize_author(m.group(1)) for m in CITE_AUTHOR_RGX.finditer(ref_text)}
            cited.update(ref_authors)

    articles.append({
        "id": art_id,
        "title": title,
        "authors": authors,
        "cited": list(cited)
    })

    # Добавляем связи соавторства
    for a1, a2 in combinations(set(authors), 2):
        key = tuple(sorted([a1, a2]))
        coauthor_edges[key] += 1

# Формируем узлы
nodes = []
author_nodes = set()

for art in articles:
    # Узел статьи
    nodes.append({
        "Id": art["id"],
        "Label": art["title"][:50],  # Обрезаем длинные названия
        "Type": "Article",
        "Authors": ", ".join(art["authors"])
    })

    # Узлы авторов
    for author in art["authors"] + art["cited"]:
        author_nodes.add(author)

for author in author_nodes:
    nodes.append({
        "Id": f"AUT_{author}",
        "Label": author,
        "Type": "Author"
    })

# Формируем ребра
edges = []
for art in articles:
    # Цитирования (с учетом частоты)
    citation_count = defaultdict(int)
    for author in art["cited"]:
        citation_count[author] += 1

    for author, count in citation_count.items():
        edges.append({
            "Source": art["id"],
            "Target": f"AUT_{author}",
            "Type": "Directed",
            "Label": "cites",
            "Weight": count
        })

    # Соавторство
for (a1, a2), count in coauthor_edges.items():
    edges.append({
        "Source": f"AUT_{a1}",
        "Target": f"AUT_{a2}",
        "Type": "Undirected",
        "Label": "coauthor",
        "Weight": count
    })

# Сохраняем данные
pd.DataFrame(nodes).to_csv(NODES_CSV, index=False)
pd.DataFrame(edges).to_csv(EDGES_CSV, index=False)
print(f"Файлы сохранены: {NODES_CSV}, {EDGES_CSV}")